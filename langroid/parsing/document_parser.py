from __future__ import annotations

import itertools
import logging
import os
import re
import tempfile
from enum import Enum
from io import BytesIO
from itertools import accumulate
from pathlib import Path
from typing import TYPE_CHECKING, Any, Dict, Generator, List, Optional, Tuple, Union

from dotenv import load_dotenv

from langroid.exceptions import LangroidImportError
from langroid.utils.object_registry import ObjectRegistry

if TYPE_CHECKING:
    import docling  # noqa
    import fitz
    import pymupdf4llm  # noqa
    import pypdf


import requests
from bs4 import BeautifulSoup

if TYPE_CHECKING:
    from PIL import Image

from langroid.mytypes import DocMetaData, Document
from langroid.parsing.parser import Parser, ParsingConfig

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    # TODO add `md` (Markdown) and `html`
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    XLSX = "xlsx"
    XLS = "xls"
    PPTX = "pptx"


def find_last_full_char(possible_unicode: bytes) -> int:
    """
    Find the index of the last full character in a byte string.
    Args:
        possible_unicode (bytes): The bytes to check.
    Returns:
        int: The index of the last full unicode character.
    """

    for i in range(len(possible_unicode) - 1, 0, -1):
        if (possible_unicode[i] & 0xC0) != 0x80:
            return i
    return 0


def is_plain_text(path_or_bytes: str | bytes) -> bool:
    """
    Check if a file is plain text by attempting to decode it as UTF-8.
    Args:
        path_or_bytes (str|bytes): The file path or bytes object.
    Returns:
        bool: True if the file is plain text, False otherwise.
    """
    if isinstance(path_or_bytes, str):
        if path_or_bytes.startswith(("http://", "https://")):
            response = requests.get(path_or_bytes)
            response.raise_for_status()
            content = response.content[:1024]
        else:
            with open(path_or_bytes, "rb") as f:
                content = f.read(1024)
    else:
        content = path_or_bytes[:1024]
    try:
        # Use magic to detect the MIME type
        import magic

        mime_type = magic.from_buffer(content, mime=True)

        # Check if the MIME type is not a text type
        if not mime_type.startswith("text/"):
            return False

        # Attempt to decode the content as UTF-8
        content = content[: find_last_full_char(content)]

        try:
            _ = content.decode("utf-8")
            # Additional checks can go here, e.g., to verify that the content
            # doesn't contain too many unusual characters for it to be considered text
            return True
        except UnicodeDecodeError:
            return False
    except UnicodeDecodeError:
        # If decoding fails, it's likely not plain text (or not encoded in UTF-8)
        return False


class DocumentParser(Parser):
    """
    Abstract base class for extracting text from special types of docs
    such as PDFs or Docx.

    Attributes:
        source (str): The source, either a URL or a file path.
        doc_bytes (BytesIO): BytesIO object containing the doc data.
    """

    @classmethod
    def create(
        cls,
        source: str | bytes,
        config: ParsingConfig,
        doc_type: str | DocumentType | None = None,
    ) -> "DocumentParser":
        """
        Create a DocumentParser instance based on source type
            and config.<source_type>.library specified.

        Args:
            source (str|bytes): The source, could be a URL, file path,
                or bytes object.
            config (ParserConfig): The parser configuration.
            doc_type (str|None): The type of document, if known

        Returns:
            DocumentParser: An instance of a DocumentParser subclass.
        """
        inferred_doc_type = DocumentParser._document_type(source, doc_type)
        if inferred_doc_type == DocumentType.PDF:
            if config.pdf.library == "fitz":
                return FitzPDFParser(source, config)
            elif config.pdf.library == "pymupdf4llm":
                return PyMuPDF4LLMParser(source, config)
            elif config.pdf.library == "docling":
                return DoclingParser(source, config)
            elif config.pdf.library == "pypdf":
                return PyPDFParser(source, config)
            elif config.pdf.library == "unstructured":
                return UnstructuredPDFParser(source, config)
            elif config.pdf.library == "pdf2image":
                return ImagePdfParser(source, config)
            elif config.pdf.library == "gemini":
                return GeminiPdfParser(source, config)
            elif config.pdf.library == "marker":
                return MarkerPdfParser(source, config)
            else:
                raise ValueError(
                    f"Unsupported PDF library specified: {config.pdf.library}"
                )
        elif inferred_doc_type == DocumentType.DOCX:
            if config.docx.library == "unstructured":
                return UnstructuredDocxParser(source, config)
            elif config.docx.library == "python-docx":
                return PythonDocxParser(source, config)
            elif config.docx.library == "markitdown-docx":
                return MarkitdownDocxParser(source, config)
            else:
                raise ValueError(
                    f"Unsupported DOCX library specified: {config.docx.library}"
                )
        elif inferred_doc_type == DocumentType.DOC:
            return UnstructuredDocParser(source, config)
        elif inferred_doc_type == DocumentType.XLS:
            return MarkitdownXLSXParser(source, config)
        elif inferred_doc_type == DocumentType.XLSX:
            return MarkitdownXLSXParser(source, config)
        elif inferred_doc_type == DocumentType.PPTX:
            return MarkitdownPPTXParser(source, config)
        else:
            source_name = source if isinstance(source, str) else "bytes"
            raise ValueError(f"Unsupported document type: {source_name}")

    def __init__(self, source: str | bytes, config: ParsingConfig):
        """
        Args:
            source (str|bytes): The source, which could be
            a path, a URL or a bytes object.
        """
        super().__init__(config)
        self.config = config
        if isinstance(source, bytes):
            self.source = "bytes"
            self.doc_bytes = BytesIO(source)
        else:
            self.source = source
            self.doc_bytes = self._load_doc_as_bytesio()

    @staticmethod
    def _document_type(
        source: str | bytes, doc_type: str | DocumentType | None = None
    ) -> DocumentType:
        """
        Determine the type of document based on the source.

        Args:
            source (str|bytes): The source, which could be a URL,
                a file path, or a bytes object.
            doc_type (str|DocumentType|None): The type of document, if known.

        Returns:
            str: The document type.
        """
        if isinstance(doc_type, DocumentType):
            return doc_type
        if doc_type:
            return DocumentType(doc_type.lower())
        if is_plain_text(source):
            return DocumentType.TXT
        if isinstance(source, str):
            # detect file type from path extension
            if source.lower().endswith(".pdf"):
                return DocumentType.PDF
            elif source.lower().endswith(".docx"):
                return DocumentType.DOCX
            elif source.lower().endswith(".doc"):
                return DocumentType.DOC
            elif source.lower().endswith(".xlsx"):
                return DocumentType.XLSX
            elif source.lower().endswith(".xls"):
                return DocumentType.XLS
            elif source.lower().endswith(".pptx"):
                return DocumentType.PPTX
            else:
                raise ValueError(f"Unsupported document type: {source}")
        else:
            # must be bytes: attempt to detect type from content
            # using magic mime type detection
            import magic

            mime_type = magic.from_buffer(source, mime=True)
            if mime_type == "application/pdf":
                return DocumentType.PDF
            elif mime_type in [
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document",
            ]:
                return DocumentType.DOCX
            elif mime_type == "application/msword":
                return DocumentType.DOC
            elif (
                mime_type
                == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ):
                return DocumentType.XLSX
            elif mime_type == "application/vnd.ms-excel":
                return DocumentType.XLS
            else:
                raise ValueError("Unsupported document type from bytes")

    def _load_doc_as_bytesio(self) -> BytesIO:
        """
        Load the docs into a BytesIO object.

        Returns:
            BytesIO: A BytesIO object containing the doc data.
        """
        if self.source.startswith(("http://", "https://")):
            response = requests.get(self.source)
            response.raise_for_status()
            return BytesIO(response.content)
        else:
            with open(self.source, "rb") as f:
                return BytesIO(f.read())

    @staticmethod
    def chunks_from_path_or_bytes(
        source: str | bytes,
        parser: Parser,
        doc_type: str | DocumentType | None = None,
        lines: int | None = None,
    ) -> List[Document]:
        """
        Get document chunks from a file path or bytes object.
        Args:
            source (str|bytes): The source, which could be a URL, path or bytes object.
            parser (Parser): The parser instance (for splitting the document).
            doc_type (str|DocumentType|None): The type of document, if known.
            lines (int|None): The number of lines to read from a plain text file.
        Returns:
            List[Document]: A list of `Document` objects,
                each containing a chunk of text, determined by the
                chunking and splitting settings in the parser config.
        """
        dtype: DocumentType = DocumentParser._document_type(source, doc_type)
        if dtype in [
            DocumentType.PDF,
            DocumentType.DOC,
            DocumentType.DOCX,
            DocumentType.PPTX,
            DocumentType.XLS,
            DocumentType.XLSX,
        ]:
            doc_parser = DocumentParser.create(
                source,
                parser.config,
                doc_type=doc_type,
            )
            chunks = doc_parser.get_doc_chunks()
            if len(chunks) == 0 and dtype == DocumentType.PDF:
                doc_parser = ImagePdfParser(source, parser.config)
                chunks = doc_parser.get_doc_chunks()
            return chunks
        else:
            # try getting as plain text; these will be chunked downstream
            # -- could be a bytes object or a path
            if isinstance(source, bytes):
                content = source.decode()
                if lines is not None:
                    file_lines = content.splitlines()[:lines]
                    content = "\n".join(line.strip() for line in file_lines)
            else:
                with open(source, "r") as f:
                    if lines is not None:
                        file_lines = list(itertools.islice(f, lines))
                        content = "\n".join(line.strip() for line in file_lines)
                    else:
                        content = f.read()
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text()
            source_name = source if isinstance(source, str) else "bytes"
            doc = Document(
                content=text,
                metadata=DocMetaData(source=str(source_name)),
            )
            return parser.split([doc])

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """Yield each page in the PDF."""
        raise NotImplementedError

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Langroid Document object (with possible metadata)
        corresponding to a given page.
        """
        raise NotImplementedError

    def fix_text(self, text: str) -> str:
        """
        Fix text extracted from a PDF.

        Args:
            text (str): The extracted text.

        Returns:
            str: The fixed text.
        """
        # Some pdf parsers introduce extra space before hyphen,
        # so use regular expression to replace 'space-hyphen' with just 'hyphen'
        return re.sub(r" +\-", "-", text)

    def get_doc(self) -> Document:
        """
        Get entire text from source as a single document.

        Returns:
            a `Document` object containing the content of the pdf file,
                and metadata containing source name (URL or path)
        """

        text = "".join(
            [
                self.get_document_from_page(page).content
                for _, page in self.iterate_pages()
            ]
        )
        return Document(content=text, metadata=DocMetaData(source=self.source))

    def get_doc_chunks(self) -> List[Document]:
        """
        Get document chunks from a pdf source,
        with page references in the document metadata.

        Returns:
            List[Document]: a list of `Document` objects,
                each containing a chunk of text
        """

        split = []  # tokens in curr split
        pages: List[str] = []
        docs: List[Document] = []
        # metadata.id to be shared by ALL chunks of this document
        common_id = ObjectRegistry.new_id()
        n_chunks = 0  # how many chunk so far
        for i, page in self.iterate_pages():
            # not used but could be useful, esp to blend the
            # metadata from the pages into the chunks
            page_doc = self.get_document_from_page(page)
            page_text = page_doc.content
            split += self.tokenizer.encode(page_text)
            pages.append(str(i + 1))
            # split could be so long it needs to be split
            # into multiple chunks. Or it could be so short
            # that it needs to be combined with the next chunk.
            while len(split) > self.config.chunk_size:
                # pretty formatting of pages (e.g. 1-3, 4, 5-7)
                p_0 = int(pages[0]) - self.config.page_number_offset
                p_n = int(pages[-1]) - self.config.page_number_offset
                page_str = f"pages {p_0}-{p_n}" if p_0 != p_n else f"page {p_0}"
                text = self.tokenizer.decode(split[: self.config.chunk_size])
                docs.append(
                    Document(
                        content=text,
                        metadata=DocMetaData(
                            source=f"{self.source} {page_str}",
                            is_chunk=True,
                            id=common_id,
                        ),
                    )
                )
                n_chunks += 1
                split = split[self.config.chunk_size - self.config.overlap :]
                pages = [str(i + 1)]
        # there may be a last split remaining:
        # if it's shorter than the overlap, we shouldn't make a chunk for it
        # since it's already included in the prior chunk;
        # the only exception is if there have been no chunks so far.
        if len(split) > self.config.overlap or n_chunks == 0:
            p_0 = int(pages[0]) - self.config.page_number_offset
            p_n = int(pages[-1]) - self.config.page_number_offset
            page_str = f"pages {p_0}-{p_n}" if p_0 != p_n else f"page {p_0}"
            text = self.tokenizer.decode(split[: self.config.chunk_size])
            docs.append(
                Document(
                    content=text,
                    metadata=DocMetaData(
                        source=f"{self.source} {page_str}",
                        is_chunk=True,
                        id=common_id,
                    ),
                )
            )
        self.add_window_ids(docs)
        return docs


class FitzPDFParser(DocumentParser):
    """
    Parser for processing PDFs using the `fitz` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, "fitz.Page"], None, None]:
        """
        Yield each page in the PDF using `fitz`.

        Returns:
            Generator[fitz.Page]: Generator yielding each page.
        """
        try:
            import fitz
        except ImportError:
            LangroidImportError("fitz", "doc-chat")
        doc = fitz.open(stream=self.doc_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            yield i, page
        doc.close()

    def get_document_from_page(self, page: "fitz.Page") -> Document:
        """
        Get Document object from a given `fitz` page.

        Args:
            page (fitz.Page): The `fitz` page object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.get_text()),
            metadata=DocMetaData(source=self.source),
        )


class PyMuPDF4LLMParser(DocumentParser):
    """
    Parser for processing PDFs using the `pymupdf4llm` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, "fitz.Page"], None, None]:
        """
        Yield each page in the PDF using `fitz`.

        Returns:
            Generator[fitz.Page]: Generator yielding each page.
        """
        try:
            import pymupdf4llm  # noqa
            import fitz
        except ImportError:
            raise LangroidImportError(
                "pymupdf4llm", ["pymupdf4llm", "all", "pdf-parsers", "doc-chat"]
            )
        doc: fitz.Document = fitz.open(stream=self.doc_bytes, filetype="pdf")
        pages: List[Dict[str, Any]] = pymupdf4llm.to_markdown(doc, page_chunks=True)
        for i, page in enumerate(pages):
            yield i, page
        doc.close()

    def get_document_from_page(self, page: Dict[str, Any]) -> Document:
        """
        Get Document object corresponding to a given "page-chunk"
        dictionary, see:
         https://pymupdf.readthedocs.io/en/latest/pymupdf4llm/api.html


        Args:
            page (Dict[str,Any]): The "page-chunk" dictionary.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.get("text", "")),
            # TODO could possible use other metadata from page, see above link.
            metadata=DocMetaData(source=self.source),
        )


class DoclingParser(DocumentParser):
    """
    Parser for processing PDFs using the `docling` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Yield each page in the PDF using `docling`.
        Code largely from this example:
        https://github.com/DS4SD/docling/blob/4d41db3f7abb86c8c65386bf94e7eb0bf22bb82b/docs/examples/export_figures.py

        Returns:
            Generator[docling.Page]: Generator yielding each page.
        """
        try:
            import docling  # noqa
        except ImportError:
            raise LangroidImportError(
                "docling", ["docling", "pdf-parsers", "all", "doc-chat"]
            )

        from docling.datamodel.base_models import InputFormat  # type: ignore
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import (  # type: ignore
            ConversionResult,
            DocumentConverter,
            PdfFormatOption,
        )
        from docling_core.types.doc import ImageRefMode  # type: ignore

        IMAGE_RESOLUTION_SCALE = 2.0
        pipeline_options = PdfPipelineOptions()
        pipeline_options.images_scale = IMAGE_RESOLUTION_SCALE
        pipeline_options.generate_page_images = True
        pipeline_options.generate_picture_images = True

        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )
        doc_path = self.source
        if doc_path == "bytes":
            # write to tmp file, then use that path
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(self.doc_bytes.getvalue())
                doc_path = temp_file.name

        output_dir = Path(str(Path(doc_path).with_suffix("")) + "-pages")
        os.makedirs(output_dir, exist_ok=True)

        result: ConversionResult = converter.convert(doc_path)

        def n_page_elements(page) -> int:  # type: ignore
            if page.assembled is None:
                return 0
            return 1 + len(page.assembled.elements)

        page_element_count = [n_page_elements(i) for i in result.pages]
        element_page_cutoff = list(accumulate([1] + page_element_count))
        for i, page in enumerate(result.pages):
            page_start = element_page_cutoff[i]
            page_end = element_page_cutoff[i + 1]
            md_file = output_dir / f"page_{i}.md"
            # we could have just directly exported to a markdown string,
            # but we need to save to a file to force generation of image-files.
            result.document.save_as_markdown(
                md_file,
                image_mode=ImageRefMode.REFERENCED,
                from_element=page_start,
                to_element=page_end,
            )
            yield i, md_file

    def get_document_from_page(self, md_file: str) -> Document:
        """
        Get Document object from a given 1-page markdown file,
        possibly containing image refs.

        Args:
            md_file (str): The markdown file path for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        with open(md_file, "r") as f:
            text = f.read()
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class PyPDFParser(DocumentParser):
    """
    Parser for processing PDFs using the `pypdf` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, pypdf.PageObject], None, None]:
        """
        Yield each page in the PDF using `pypdf`.

        Returns:
            Generator[pypdf.pdf.PageObject]: Generator yielding each page.
        """
        try:
            import pypdf
        except ImportError:
            raise LangroidImportError("pypdf", "pdf-parsers")
        reader = pypdf.PdfReader(self.doc_bytes)
        for i, page in enumerate(reader.pages):
            yield i, page

    def get_document_from_page(self, page: pypdf.PageObject) -> Document:
        """
        Get Document object from a given `pypdf` page.

        Args:
            page (pypdf.pdf.PageObject): The `pypdf` page object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page.extract_text()),
            metadata=DocMetaData(source=self.source),
        )


class ImagePdfParser(DocumentParser):
    """
    Parser for processing PDFs that are images, i.e. not "true" PDFs.
    """

    def iterate_pages(
        self,
    ) -> Generator[Tuple[int, "Image"], None, None]:  # type: ignore
        try:
            from pdf2image import convert_from_bytes
        except ImportError:
            raise LangroidImportError("pdf2image", "pdf-parsers")

        images = convert_from_bytes(self.doc_bytes.getvalue())
        for i, image in enumerate(images):
            yield i, image

    def get_document_from_page(self, page: "Image") -> Document:  # type: ignore
        """
        Get Document object corresponding to a given `pdf2image` page.

        Args:
            page (Image): The PIL Image object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        try:
            import pytesseract
        except ImportError:
            raise LangroidImportError("pytesseract", "pdf-parsers")

        text = pytesseract.image_to_string(page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredPDFParser(DocumentParser):
    """
    Parser for processing PDF files using the `unstructured` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.pdf import partition_pdf
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        # from unstructured.chunking.title import chunk_by_title

        try:
            elements = partition_pdf(file=self.doc_bytes, include_page_breaks=True)
        except Exception as e:
            raise Exception(
                f"""
                Error parsing PDF: {e}
                The `unstructured` library failed to parse the pdf.
                Please try a different library by setting the `library` field
                in the `pdf` section of the `parsing` field in the config file.
                Other supported libraries are:
                fitz, pymupdf4llm, pypdf
                """
            )

        # elements = chunk_by_title(elements)
        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given `unstructured` element.

        Args:
            page (unstructured element): The `unstructured` element object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        text = " ".join(el.text for el in page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredDocxParser(DocumentParser):
    """
    Parser for processing DOCX files using the `unstructured` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.docx import partition_docx
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        elements = partition_docx(file=self.doc_bytes, include_page_breaks=True)

        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given `unstructured` element.

        Note:
            The concept of "pages" doesn't actually exist in the .docx file format in
            the same way it does in formats like .pdf. A .docx file is made up of a
            series of elements like paragraphs and tables, but the division into
            pages is done dynamically based on the rendering settings (like the page
            size, margin size, font size, etc.).

        Args:
            page (unstructured element): The `unstructured` element object.

        Returns:
            Document object, with content and possible metadata.
        """
        text = " ".join(el.text for el in page)
        return Document(
            content=self.fix_text(text),
            metadata=DocMetaData(source=self.source),
        )


class UnstructuredDocParser(UnstructuredDocxParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:  # type: ignore
        try:
            from unstructured.partition.doc import partition_doc
        except ImportError:
            raise ImportError(
                """
                The `unstructured` library is not installed by default with langroid.
                To include this library, please install langroid with the
                `unstructured` extra by running `pip install "langroid[unstructured]"`
                or equivalent.
                """
            )

        elements = partition_doc(file=self.doc_bytes, include_page_breaks=True)

        page_number = 1
        page_elements = []  # type: ignore
        for el in elements:
            if el.category == "PageBreak":
                if page_elements:  # Avoid yielding empty pages at the start
                    yield page_number, page_elements
                page_number += 1
                page_elements = []
            else:
                page_elements.append(el)
        # Yield the last page if it's not empty
        if page_elements:
            yield page_number, page_elements


class PythonDocxParser(DocumentParser):
    """
    Parser for processing DOCX files using the `python-docx` library.
    """

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Simulate iterating through pages.
        In a DOCX file, pages are not explicitly defined,
        so we consider each paragraph as a separate 'page' for simplicity.
        """
        try:
            import docx
        except ImportError:
            raise LangroidImportError("python-docx", "docx")

        doc = docx.Document(self.doc_bytes)
        for i, para in enumerate(doc.paragraphs, start=1):
            yield i, [para]

    def get_document_from_page(self, page: Any) -> Document:
        """
        Get Document object from a given 'page', which in this case is a single
        paragraph.

        Args:
            page (list): A list containing a single Paragraph object.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        paragraph = page[0]
        return Document(
            content=self.fix_text(paragraph.text),
            metadata=DocMetaData(source=self.source),
        )


class MarkitdownDocxParser(DocumentParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        try:
            from markitdown import MarkItDown
        except ImportError:
            LangroidImportError("markitdown", ["markitdown", "doc-parsers"])
        md = MarkItDown()
        self.doc_bytes.seek(0)  # Reset to start

        # Direct conversion from stream works for DOCX (unlike XLSX)
        result = md.convert_stream(self.doc_bytes, file_extension=".docx")

        # Split content into logical sections (paragraphs, sections, etc.)
        # This approach differs from the strict page-based approach used for PDFs
        sections = re.split(r"(?=# |\n## |\n### )", result.text_content)

        # Filter out empty sections
        sections = [section for section in sections if section.strip()]

        for i, section in enumerate(sections):
            yield i, section

    def get_document_from_page(self, md_content: str) -> Document:
        """
        Get Document object from a given markdown section.

        Args:
            md_content (str): The markdown content for the section.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(md_content),
            metadata=DocMetaData(source=self.source),
        )


class MarkitdownXLSXParser(DocumentParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        try:
            from markitdown import MarkItDown
        except ImportError:
            LangroidImportError("markitdown", "doc-parsers")
        md = MarkItDown()
        self.doc_bytes.seek(0)  # Reset to start

        # Save stream to a temp file since md.convert() expects a path or URL
        # Temporary workaround until markitdown fixes convert_stream function
        # for xls and xlsx files
        # See issue here https://github.com/microsoft/markitdown/issues/321
        with tempfile.NamedTemporaryFile(delete=True, suffix=".xlsx") as temp_file:
            temp_file.write(self.doc_bytes.read())
            temp_file.flush()  # Ensure data is written before reading
            result = md.convert(temp_file.name)

        sheets = re.split(r"(?=## Sheet\d+)", result.text_content)

        for i, sheet in enumerate(sheets):
            yield i, sheet

    def get_document_from_page(self, md_content: str) -> Document:
        """
        Get Document object from a given 1-page markdown string.

        Args:
            md_content (str): The markdown content for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(md_content),
            metadata=DocMetaData(source=self.source),
        )


class MarkitdownPPTXParser(DocumentParser):
    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        try:
            from markitdown import MarkItDown
        except ImportError:
            LangroidImportError("markitdown", "doc-parsers")

        md = MarkItDown()
        self.doc_bytes.seek(0)
        result = md.convert_stream(self.doc_bytes, file_extension=".pptx")
        slides = re.split(r"(?=<!-- Slide number: \d+ -->)", result.text_content)
        for i, slide in enumerate(slides):
            yield i, slide

    def get_document_from_page(self, md_content: str) -> Document:
        """
        Get Document object from a given 1-page markdown string.

        Args:
            md_content (str): The markdown content for the page.

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(md_content),
            metadata=DocMetaData(source=self.source),
        )


class GeminiPdfParser(DocumentParser):
    """
    This class converts PDFs to Markdown using Gemini multimodal LLMs.

    It extracts pages, converts them with the LLM (replacing images with
    detailed descriptions), and outputs Markdown page by page. The
    conversion follows `GEMINI_SYSTEM_INSTRUCTION`. It employs
    multiprocessing for speed, async requests with rate limiting, and
    handles errors.

    It supports page-by-page splitting or chunking multiple pages into
    one, respecting page boundaries and a `max_token_limit`.
    """

    DEFAULT_MAX_TOKENS = 7000
    OUTPUT_DIR = Path(".gemini_pdfparser")  # Fixed output directory

    GEMINI_SYSTEM_INSTRUCTION = """
    ### **Convert PDF to Markdown**
    1. **Text:**
        * Preserve structure, formatting (**bold**, *italic*), lists, and indentation.
        * **Remove running heads (page numbers, headers/footers).**
        * Keep section and chapter titles; discard repeated page headers.
    2. **Images:** Replace with **detailed, creative descriptions**
    optimized for clarity and understanding.
    3. **Tables:** Convert to Markdown tables with proper structure.
    4. **Math:** Use LaTeX (`...` inline, `$...$` block).
    5. **Code:** Wrap in fenced blocks without specifying a language:

        ```
        code
        ```
    6. **Clean Output:**
        * No system messages, metadata, or artifacts or ```markdown``` identifier.
        * Do **not** include introductory or explanatory messages
        like "Here is your output."
        * Ensure formatting is **consistent and structured**
        for feeding into a markdown parser.
    """.strip()

    def __init__(self, source: Union[str, bytes], config: ParsingConfig):
        super().__init__(source, config)
        if not config.pdf.gemini_config:
            raise ValueError(
                "GeminiPdfParser requires a Gemini-based config in pdf parsing config"
            )
        self.model_name = config.pdf.gemini_config.model_name

        # Ensure output directory exists
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        prefix = (
            Path(source).stem + "_"
            if isinstance(source, str) and Path(source).exists()
            else "output_"
        )
        temp_file = tempfile.NamedTemporaryFile(
            suffix=".md",
            prefix=prefix,
            dir=str(self.OUTPUT_DIR),
            delete=False,
        )
        temp_file.close()
        self.output_filename = Path(temp_file.name)

        self.max_tokens = config.pdf.gemini_config.max_tokens or self.DEFAULT_MAX_TOKENS

        """
        If True, each PDF page is processed as a separate chunk,
        resulting in one LLM request per page. If False, pages are
        grouped into chunks based on `max_token_limit` before being sent
        to the LLM.
        """
        self.split_on_page = config.pdf.gemini_config.split_on_page or False

        # Rate limiting parameters
        import asyncio

        self.requests_per_minute = config.pdf.gemini_config.requests_per_minute or 5

        """
        A semaphore to control the number of concurrent requests to the LLM,
        preventing rate limit errors.  A semaphore slot is acquired before
        making an LLM request and released after the request is complete.
        """
        self.semaphore = asyncio.Semaphore(self.requests_per_minute)
        self.retry_delay = 5  # seconds, for exponential backoff
        self.max_retries = 3

    def _extract_page(self, page_num: int) -> Dict[str, Any]:
        """
        Extracts a single page and estimates token count.
        Opens the PDF from self.doc_bytes (a BytesIO object).
        """
        import fitz

        try:
            # Always open the document from in-memory bytes.
            doc = fitz.open(stream=self.doc_bytes.getvalue(), filetype="pdf")
            new_pdf = fitz.open()
            new_pdf.insert_pdf(doc, from_page=page_num, to_page=page_num)
            pdf_bytes = new_pdf.write()
            text = doc[page_num].get_text("text")
            token_count = len(text) // 4 if text else len(pdf_bytes) // 4

            return {
                "page_numbers": page_num + 1,
                "pdf_bytes": pdf_bytes,
                "token_count": token_count,
            }
        except Exception as e:
            raise ValueError(f"Error processing PDF document: {e}") from e

    def _extract_pdf_pages_parallel(
        self, num_workers: Optional[int] = None
    ) -> List[Dict[str, Any]]:
        """Parallel PDF page extraction using self.doc_bytes."""
        from multiprocessing import Pool, cpu_count

        import fitz
        from tqdm import tqdm

        try:
            doc = fitz.open(stream=self.doc_bytes.getvalue(), filetype="pdf")
            total_pages = len(doc)
        except Exception as e:
            raise ValueError(f"Error opening PDF document: {e}") from e

        num_workers = num_workers or cpu_count()
        with Pool(num_workers) as pool:
            with tqdm(total=total_pages, desc="Extracting pages", unit="page") as pbar:
                results = []
                for result in pool.imap(self._extract_page, range(total_pages)):
                    results.append(result)
                    pbar.update(1)

        return results

    def _group_pages_by_token_limit(
        self, pages: List[Dict[str, Any]], max_tokens: int = DEFAULT_MAX_TOKENS
    ) -> List[List[Dict[str, Any]]]:
        """Groups pages into chunks where each chunk is approximately `max_tokens`."""
        chunks: List[List[Dict[str, Any]]] = []
        current_chunk: List[Dict[str, Any]] = []
        current_tokens = 0

        for page in pages:
            if current_tokens + page["token_count"] > max_tokens and current_chunk:
                chunks.append(current_chunk)
                current_chunk = []
                current_tokens = 0

            current_chunk.append(page)
            current_tokens += page["token_count"]

        if current_chunk:  # Add remaining pages
            chunks.append(current_chunk)

        return chunks

    def _merge_pages_into_pdf_with_metadata(
        self, page_group: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """
        Merges grouped pages into a single binary chunk so that
        it does not exceed max token limit
        """
        import fitz

        merged_pdf = fitz.open()
        page_numbers = []

        for page in page_group:
            temp_pdf = fitz.open("pdf", page["pdf_bytes"])
            merged_pdf.insert_pdf(temp_pdf)
            page_numbers.append(page["page_numbers"])

        return {
            "pdf_bytes": merged_pdf.write(),  # Binary PDF data
            "page_numbers": page_numbers,  # List of page numbers in this chunk
        }

    def _prepare_pdf_chunks_for_gemini(
        self,
        num_workers: Optional[int] = None,
        max_tokens: int = DEFAULT_MAX_TOKENS,
        split_on_page: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        Extracts, groups, and merges PDF pages into chunks with embedded page markers.
        """
        from multiprocessing import Pool

        pages = self._extract_pdf_pages_parallel(num_workers)

        if split_on_page:
            # Each page becomes its own chunk
            return pages
        else:
            # Group pages based on token limit
            chunks = self._group_pages_by_token_limit(pages, max_tokens)
            with Pool(num_workers) as pool:
                pdf_chunks = pool.map(self._merge_pages_into_pdf_with_metadata, chunks)
            return pdf_chunks

    async def _send_chunk_to_gemini(
        self, chunk: Dict[str, Any], gemini_api_key: str
    ) -> str:
        """
        Sends a PDF chunk to the Gemini API and returns the response text.
        Uses retries with exponential backoff to handle transient failures.
        """
        import asyncio
        import logging

        from google import genai
        from google.genai import types

        async with self.semaphore:  # Limit concurrent API requests
            for attempt in range(self.max_retries):
                try:
                    client = genai.Client(api_key=gemini_api_key)

                    # Send the request with PDF content and system instructions
                    response = await client.aio.models.generate_content(
                        model=self.model_name,
                        contents=[
                            types.Part.from_bytes(
                                data=chunk["pdf_bytes"], mime_type="application/pdf"
                            ),
                            self.GEMINI_SYSTEM_INSTRUCTION,
                        ],
                    )

                    # Return extracted text if available
                    return str(response.text) if response.text else ""

                except Exception as e:
                    # Log error with page numbers for debugging
                    logging.error(
                        "Attempt %d failed for pages %s: %s",
                        attempt + 1,
                        chunk.get("page_numbers", "Unknown"),
                        e,
                    )

                    if attempt < self.max_retries - 1:
                        # Apply exponential backoff before retrying
                        delay = self.retry_delay * (2**attempt)
                        logging.info("Retrying in %s sec...", delay)
                        await asyncio.sleep(delay)
                    else:
                        # Log failure after max retries
                        logging.error(
                            "Max retries reached for pages %s",
                            chunk.get("page_numbers", "Unknown"),
                        )
                        break

        return ""  # Return empty string if all retries fail

    async def process_chunks(
        self, chunks: List[Dict[str, Any]], api_key: str
    ) -> List[str]:
        """
        Processes PDF chunks by sending them to the Gemini API and
        collecting the results.

        Args:
            chunks: A list of dictionaries, where each dictionary represents
                a PDF chunk and contains the PDF data and page numbers.
            api_key: The Gemini API key.
        """
        # To show nice progress bar
        from tqdm.asyncio import tqdm_asyncio

        # Create a list of asynchronous tasks to send each chunk to Gemini.
        # Chunk in this case might be single page or group of pages returned
        # by prepare_pdf_chunks function
        tasks = [self._send_chunk_to_gemini(chunk, api_key) for chunk in chunks]

        # Gather the results from all tasks, allowing exceptions to be returned.
        # tqdm_asyncio is wrapper around asyncio.gather
        gathered_results = await tqdm_asyncio.gather(
            *tasks, desc="Processing chunks(pages)", unit="chunk"
        )
        results = []
        for i, result in enumerate(gathered_results):
            chunk = chunks[i]  # Get the corresponding chunk.

            if isinstance(result, Exception):
                # Handle exceptions that occurred during chunk processing.
                logging.error(
                    "Failed to process chunk %s: %s",
                    chunk.get("page_numbers", "Unknown"),
                    result,
                )
                results.append(
                    "<!----Error: Could not process chunk %s---->"
                    % chunk.get("page_numbers", "Unknown")
                )
            else:
                # Process successful results and append page/chunk markers.
                markdown = str(result)
                if self.split_on_page:
                    results.append(
                        markdown + f"<!----Page-{chunk['page_numbers']}---->"
                    )
                else:
                    results.append(
                        markdown + f"<!----Chunk-{chunk['page_numbers']}---->"
                    )

        return results  # Return the list of results.

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Iterates over the document pages, extracting content using the
        Gemini API, saves them to a markdown file, and yields page numbers
        along with their corresponding content.

        Yields:
            A generator of tuples, where each tuple contains the page number
            (int) and the page content (Any).
        """
        import asyncio
        import os

        # Load environment variables (e.g., GEMINI_API_KEY) from a .env file.
        load_dotenv()
        gemini_api_key = os.getenv("GEMINI_API_KEY")
        if not gemini_api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables.")

        try:
            # This involves extracting pages, grouping them according to the
            # `max_tokens` limit (if `split_on_page` is False), and
            # merging pages into larger PDF chunks. The result
            # is a list of dictionaries, where each dictionary contains the
            # PDF bytes and the associated page numbers or single page if
            # `split_on_page` is true

            pdf_chunks = self._prepare_pdf_chunks_for_gemini(
                num_workers=8,
                max_tokens=self.max_tokens,
                split_on_page=self.split_on_page,
            )

            # We asynchronously processes each chunk, sending it
            # to Gemini and retrieving the Markdown output. It handles rate
            # limiting and retries.
            markdown_results = asyncio.run(
                self.process_chunks(pdf_chunks, gemini_api_key)
            )

            # This file serves as an intermediate storage location for the
            # complete Markdown output.
            with open(self.output_filename, "w", encoding="utf-8") as outfile:
                outfile.write("\n\n".join(markdown_results))

            # Read the full Markdown content from the temporary file.
            with open(self.output_filename, "r", encoding="utf-8") as infile:
                full_markdown = infile.read()

            # The splitting is based on the `split_on_page` setting. If True,
            # the Markdown is split using the "Page-" marker. Otherwise, it's
            # split using the "Chunk-" marker.
            if self.split_on_page:
                pages = full_markdown.split("<!----Page-")
            else:
                pages = full_markdown.split("<!----Chunk-")

            # Remove the first element if it's empty (due to the split).
            if pages and pages[0] == "":
                pages = pages[1:]

            # Iterate over the pages or chunks and yield their content.
            for i, page in enumerate(pages):
                # Check for errors during processing.
                if "<!----Error:" in page:
                    page_content = page
                    logging.warning(f"Page {i}: Error processing chunk.")
                else:
                    # Extract the actual page content by removing the marker.
                    page_content = (
                        page.split("---->", 1)[1]
                        if len(page.split("---->", 1)) > 1
                        else page
                    )

                # Yield the page number and content.
                yield i, page_content

        except Exception as e:
            raise ValueError(f"Error processing document: {e}") from e

    def get_document_from_page(self, page: str) -> Document:
        """
        Get a Document object from a given markdown page.
        """
        return Document(
            content=page,
            metadata=DocMetaData(source=self.source),
        )


class MarkerPdfParser(DocumentParser):
    """
    Parse PDF files using the `marker` library: https://github.com/VikParuchuri/marker
    """

    DEFAULT_CONFIG = {"paginate_output": True, "output_format": "markdown"}

    def __init__(self, source: Union[str, bytes], config: ParsingConfig):
        super().__init__(source, config)
        user_config = (
            config.pdf.marker_config.config_dict if config.pdf.marker_config else {}
        )

        self.config_dict = {**MarkerPdfParser.DEFAULT_CONFIG, **user_config}

    def iterate_pages(self) -> Generator[Tuple[int, Any], None, None]:
        """
        Yield each page in the PDF using `marker`.
        """
        try:
            import marker  # noqa
        except ImportError:
            raise LangroidImportError(
                "marker-pdf", ["marker-pdf", "pdf-parsers", "all", "doc-chat"]
            )

        import re

        from marker.config.parser import ConfigParser
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import save_output

        config_parser = ConfigParser(self.config_dict)
        converter = PdfConverter(
            config=config_parser.generate_config_dict(),
            artifact_dict=create_model_dict(),
            processor_list=config_parser.get_processors(),
            renderer=config_parser.get_renderer(),
            llm_service=config_parser.get_llm_service(),
        )
        doc_path = self.source
        if doc_path == "bytes":
            # write to tmp file, then use that path
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(self.doc_bytes.getvalue())
                doc_path = temp_file.name

        output_dir = Path(str(Path(doc_path).with_suffix("")) + "-pages")
        os.makedirs(output_dir, exist_ok=True)
        filename = Path(doc_path).stem + "_converted"

        rendered = converter(doc_path)
        save_output(rendered, output_dir=output_dir, fname_base=filename)
        file_path = output_dir / f"{filename}.md"

        with open(file_path, "r", encoding="utf-8") as f:
            full_markdown = f.read()

        # Regex for splitting pages
        pages = re.split(r"\{\d+\}----+", full_markdown)

        page_no = 0
        for page in pages:
            if page.strip():
                yield page_no, page
            page_no += 1

    def get_document_from_page(self, page: str) -> Document:
        """
        Get Document object from a given 1-page markdown file,
        possibly containing image refs.

        Args:
            page (str): The page we get by splitting large md file from
            marker

        Returns:
            Document: Document object, with content and possible metadata.
        """
        return Document(
            content=self.fix_text(page),
            metadata=DocMetaData(source=self.source),
        )
